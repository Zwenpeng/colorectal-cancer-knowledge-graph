#!/usr/bin/env python
import argparse
import base64
import copy
import csv
import hashlib
import html
import json
import mimetypes
import os
import re
import shutil
import sys
import time
import subprocess
import tempfile
import urllib.error
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape


BASE_KG_DIR = Path("colorectal_knowledge_graph")
DEFAULT_OUT_DIR = Path("kg_update_system") / "runs"
SUPPORTED_EXTS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".html",
    ".htm",
    ".docx",
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tif",
    ".tiff",
    ".webp",
}

TEXT_EXTS = {".txt", ".md", ".csv", ".json", ".html", ".htm"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}

NODE_FIELDS = [
    "id",
    "label",
    "node_kind",
    "in_colorectal_tree",
    "category",
    "research_axis",
    "semantic_types",
    "depth_from_root",
    "active",
    "concept_status",
    "umls_cui",
    "icd_o_3_code",
    "definition",
    "aliases",
    "alias_count",
    "nci_url",
    "kg_status",
    "source_docs",
    "evidence",
]

EDGE_FIELDS = [
    "source",
    "source_label",
    "relation_code",
    "relation_type",
    "relation_label_cn",
    "target",
    "target_label",
    "relation_group",
    "direction",
    "research_tier",
    "source_docs",
    "evidence",
]

RELATION_LABELS = {
    "MENTIONED_WITH": "同文共现",
    "LLM_RELATED_TO": "模型抽取相关",
    "LLM_ASSOCIATED_WITH": "模型抽取关联",
    "LLM_CAUSES_OR_DRIVES": "可能驱动/导致",
    "LLM_TREATS": "治疗/用于",
    "LLM_BIOMARKER_OF": "生物标志物",
    "LLM_LOCATED_IN": "位于/发生于",
    "LLM_SUBTYPE_OF": "亚型属于",
}

CATEGORY_TO_AXIS = {
    "Disease Concept": "Disease taxonomy",
    "Histology": "Pathology/cytology",
    "Stage/Grade": "Clinical stage/course",
    "Clinical Course": "Clinical stage/course",
    "Molecular Disease Subtype": "Molecular genetics",
    "Gene/Genome": "Molecular genetics",
    "Molecular Abnormality": "Molecular genetics",
    "Protein/Gene Product": "Molecular genetics",
    "Anatomic Site": "Anatomy",
    "Cell": "Pathology/cytology",
    "Tissue": "Pathology/cytology",
    "Finding/Clinical Attribute": "Phenotype/finding",
    "Treatment Regimen": "Therapy",
    "Associated Disease/Neoplasm": "Associated conditions",
    "Other Entity": "Other",
}


def read_csv(path):
    with Path(path).open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def write_csv(path, rows, fields):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with Path(path).open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fields})


def clean_text(text):
    text = re.sub(r"<script[\s\S]*?</script>", " ", str(text), flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def strip_markdown(text):
    text = re.sub(r"```[\s\S]*?```", " ", str(text))
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", " ", text)
    text = re.sub(r"\[[^\]]+\]\([^)]+\)", r"\1", text)
    text = re.sub(r"[#>*_`~\-]{1,}", " ", text)
    return clean_text(text)


def guess_title(text, fallback):
    text = clean_text(text)
    if not text:
        return fallback
    lines = [line.strip() for line in str(text).splitlines() if line.strip()]
    if lines:
        first = clean_text(lines[0])[:120].strip()
        if 4 <= len(first) <= 120:
            return first
    return fallback


def ensure_python_package_path():
    candidates = []
    env_runtime = os.environ.get("CODEX_WORKSPACE_PYTHON", "")
    if env_runtime:
        candidates.append(Path(env_runtime).expanduser())
    candidates.append(Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python")
    candidates.append(Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python" / "Lib" / "site-packages")
    for runtime in candidates:
        if runtime.exists():
            runtime_str = str(runtime)
            if runtime_str not in sys.path:
                sys.path.insert(0, runtime_str)


def decode_bytes(data):
    if not data:
        return ""
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def find_tesseract():
    candidates = [
        shutil.which("tesseract"),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    return [str(Path(candidate)) for candidate in candidates if candidate and Path(candidate).exists()]


def find_tessdata_dir():
    local_tessdata = Path(__file__).resolve().parent / "tessdata"
    if (local_tessdata / "chi_sim.traineddata").exists() and (local_tessdata / "eng.traineddata").exists():
        return str(local_tessdata)
    program_tessdata = Path(r"C:\Program Files\Tesseract-OCR\tessdata")
    if (program_tessdata / "chi_sim.traineddata").exists() and (program_tessdata / "eng.traineddata").exists():
        return str(program_tessdata)
    return ""


def find_poppler_tool(name):
    candidates = [
        shutil.which(name),
        Path.home()
        / "AppData"
        / "Local"
        / "Microsoft"
        / "WinGet"
        / "Packages"
        / "oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
        / "poppler-25.07.0"
        / "Library"
        / "bin"
        / f"{name}.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return ""


def image_to_markdown_via_llm(path, config):
    if not config or not config.get("llm", {}).get("enabled", False):
        return "", "未启用大模型视觉读取。"
    mime = mimetypes.guess_type(str(path))[0] or "image/png"
    encoded = base64.b64encode(Path(path).read_bytes()).decode("ascii")
    messages = [
        {
            "role": "system",
            "content": "你是医学资料图片读取助手。只根据图片内容提取文字、表格、标题和医学实体，不要编造图片外信息。",
        },
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "请读取这张图片，输出可继续做知识图谱抽取的 Markdown。"
                        "若图片不是医学资料，也要如实描述可见文字；若完全无法读取，请说明原因。"
                    ),
                },
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{encoded}"}},
            ],
        },
    ]
    try:
        content = call_chat_api(config, messages, json_mode=False, max_tokens=3000)
    except Exception as exc:
        return "", f"大模型图片读取失败：{exc}"
    text = clean_text(content)
    if not text:
        return "", "大模型未从图片返回可处理文本。"
    return text, ""


def pdf_to_markdown_via_llm(path, config, max_pages=3):
    pdftoppm = find_poppler_tool("pdftoppm")
    if not pdftoppm:
        return "", "扫描版 PDF 需要 Poppler 的 pdftoppm 才能转图片后交给大模型读取。"
    texts = []
    errors = []
    with tempfile.TemporaryDirectory(prefix="kg_pdf_pages_") as tmp:
        prefix = str(Path(tmp) / "page")
        try:
            output = subprocess.run(
                [pdftoppm, "-png", "-f", "1", "-l", str(max_pages), str(Path(path).resolve()), prefix],
                capture_output=True,
                text=False,
                check=False,
                timeout=180,
            )
        except Exception as exc:
            return "", f"PDF 转图片失败：{exc}"
        stderr = clean_text(decode_bytes(output.stderr))
        if stderr:
            errors.append(stderr)
        images = sorted(Path(tmp).glob("page-*.png"))
        if not images:
            return "", "PDF 转图片后未生成页面图片。" + (f" {stderr}" if stderr else "")
        for idx, image_path in enumerate(images, start=1):
            text, error = image_to_markdown_via_llm(image_path, config)
            if text:
                texts.append(f"## PDF第{idx}页图像读取\n{text}")
            if error:
                errors.append(f"第{idx}页：{error}")
    if texts:
        note = "扫描版 PDF 已转为图片并由大模型读取；请人工核对。" + (f" 备注：{'；'.join(errors[:3])}" if errors else "")
        return clean_text("\n\n".join(texts)), note
    return "", "；".join(errors) or "大模型未能从扫描版 PDF 页面读取文本。"


def ocr_image(path):
    ensure_python_package_path()
    errors = []
    tessdata_dir = find_tessdata_dir()
    for candidate in find_tesseract():
        try:
            cmd = [candidate, str(Path(path).resolve()), "stdout", "--psm", "6", "-l", "chi_sim+eng"]
            if tessdata_dir:
                cmd[3:3] = ["--tessdata-dir", tessdata_dir]
            output = subprocess.run(
                cmd,
                capture_output=True,
                text=False,
                check=False,
                timeout=120,
            )
            text = clean_text(decode_bytes(output.stdout))
            if text:
                return text, ""
            stderr = clean_text(decode_bytes(output.stderr))
            if stderr:
                errors.append(stderr)
        except Exception as exc:
            errors.append(str(exc))
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(path)
        config = f'--tessdata-dir "{tessdata_dir}"' if tessdata_dir else ""
        text = pytesseract.image_to_string(image, lang="chi_sim+eng", config=config)
        text = clean_text(text)
        if text:
            return text, ""
    except Exception as exc:
        errors.append(f"Pillow/pytesseract 不可用：{exc}")
    if errors:
        return "", "图片文件已收到，但 OCR 未能提取文字；" + "；".join(errors[:3])
    return "", "图片文件已收到，但当前电脑未检测到 OCR 引擎 Tesseract；请启用大模型增强读取图片，或先把图片文字粘贴为文本。"


def extract_docx_text(path):
    ensure_python_package_path()
    try:
        from docx import Document
    except Exception as exc:
        return "", f"python-docx unavailable: {exc}"
    doc = Document(str(path))
    paras = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cell_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if cell_text:
                paras.append(cell_text)
    return clean_text("\n".join(paras)), ""


def extract_pdf_text(path):
    ensure_python_package_path()
    text_parts = []
    errors = []
    try:
        import fitz
        doc = fitz.open(str(path))
        for page in doc:
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)
        doc.close()
    except Exception as exc:
        errors.append(f"PyMuPDF：{exc}")
    text = clean_text("\n".join(text_parts))
    if text:
        return text, "；".join(errors)
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    except Exception as exc:
        errors.append(f"pypdf：{exc}")
    text = clean_text("\n".join(text_parts))
    if text:
        return text, "；".join(errors)
    try:
        poppler = find_poppler_tool("pdftotext")
        for poppler in [p for p in [poppler] if p and Path(p).exists()]:
            out = subprocess.run(
                [str(poppler), "-layout", str(Path(path).resolve()), "-"],
                capture_output=True,
                text=False,
                check=False,
                timeout=120,
            )
            text = clean_text(decode_bytes(out.stdout))
            stderr = clean_text(decode_bytes(out.stderr))
            if stderr:
                errors.append(f"pdftotext：{stderr}")
            if text:
                return text, "；".join(errors)
    except Exception as exc:
        errors.append(f"pdftotext：{exc}")
    if not errors:
        errors.append("未检测到可用 PDF 解析器。")
    errors.append("PDF 未提取到文本；如果这是扫描版 PDF，需要 OCR 或启用支持图片/PDF视觉的大模型。")
    return "", "；".join(errors)


def extract_file_text(path, config=None, use_llm=False):
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTS:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="gb18030", errors="ignore")
        if suffix == ".csv":
            text = re.sub(r"[,;\t]", " ", text)
        elif suffix in {".html", ".htm"}:
            text = clean_text(text)
        elif suffix == ".json":
            try:
                text = json.dumps(json.loads(text), ensure_ascii=False)
            except json.JSONDecodeError:
                pass
        return clean_text(text), ""
    if suffix in DOCX_EXTS:
        return extract_docx_text(path)
    if suffix in PDF_EXTS:
        text, error = extract_pdf_text(path)
        if text:
            return text, error
        if use_llm and config and config.get("llm", {}).get("enabled", False):
            llm_text, llm_error = pdf_to_markdown_via_llm(path, config)
            if llm_text:
                return llm_text, "本 PDF 未提取到内嵌文本，已尝试转图片并由大模型读取；请人工核对。"
            error = "；".join(part for part in [error, llm_error] if part)
        return "", error
    if suffix in IMAGE_EXTS:
        text, error = ocr_image(path)
        if text:
            return text, error
        if use_llm and config and config.get("llm", {}).get("enabled", False):
            llm_text, llm_error = image_to_markdown_via_llm(path, config)
            if llm_text:
                note = "本图片文字由大模型视觉读取生成；请人工核对。" + (f" OCR备注：{error}" if error else "")
                return llm_text, note
            error = "；".join(part for part in [error, llm_error] if part)
        return "", error
    return "", f"unsupported file type: {suffix}"


def normalize(text):
    return re.sub(r"\s+", " ", str(text).strip().lower())


def stable_new_id(label):
    digest = hashlib.sha1(normalize(label).encode("utf-8")).hexdigest()[:12].upper()
    return f"NEW_{digest}"


def snippet(text, start, end, radius=110):
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    prefix = "..." if left > 0 else ""
    suffix = "..." if right < len(text) else ""
    return prefix + text[left:right].strip() + suffix


def load_documents(input_path, config=None, use_llm=False, return_errors=False):
    path = Path(input_path)
    files = []
    parse_errors = []
    if not path.exists():
        parse_errors.append({"path": str(path), "error": "路径不存在。"})
    elif path.is_file():
        files = [path]
    else:
        files = sorted(p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS)
        if not files:
            parse_errors.append({"path": str(path), "error": f"目录中未发现支持格式文件。支持格式：{', '.join(sorted(SUPPORTED_EXTS))}"})
    docs = []
    for file in files:
        try:
            text, error = extract_file_text(file, config=config, use_llm=use_llm)
        except Exception as exc:
            text, error = "", str(exc)
        if text:
            docs.append(
                {
                    "doc_id": file.stem,
                    "path": str(file.resolve()),
                    "text": text,
                    "title": guess_title(text, file.stem),
                    "file_type": file.suffix.lower().lstrip("."),
                    "parse_error": error,
                }
            )
        else:
            parse_errors.append(
                {
                    "path": str(file.resolve()),
                    "file_type": file.suffix.lower().lstrip("."),
                    "error": error or "未提取到可处理文本。",
                }
            )
    if return_errors:
        return docs, parse_errors
    return docs


def format_parse_errors(parse_errors, limit=12):
    if not parse_errors:
        return ""
    lines = []
    for item in parse_errors[:limit]:
        name = Path(item.get("path", "")).name or item.get("path", "")
        file_type = item.get("file_type", "")
        prefix = f"{name}" + (f" ({file_type})" if file_type else "")
        lines.append(f"- {prefix}: {item.get('error', '')}")
    if len(parse_errors) > limit:
        lines.append(f"- 另有 {len(parse_errors) - limit} 个文件未列出。")
    return "\n".join(lines)


def load_base_kg(kg_dir):
    kg_dir = Path(kg_dir)
    nodes = {row["id"]: row for row in read_csv(kg_dir / "kg_nodes.csv")}
    edge_path = kg_dir / "kg_edges_research_core.csv"
    if not edge_path.exists():
        edge_path = kg_dir / "kg_edges.csv"
    edges = read_csv(edge_path)
    return nodes, edges


def build_lexicon(nodes):
    entries = []
    for node in nodes.values():
        terms = [node.get("label", "")]
        aliases = node.get("aliases", "")
        if aliases:
            terms.extend(aliases.split("|"))
        for term in terms:
            term = re.sub(r"\s+", " ", term).strip()
            if len(term) < 3:
                continue
            lowered = term.lower()
            if lowered in {"cancer", "tumor", "stage", "gene", "cell", "tissue"}:
                continue
            entries.append((term, lowered, node["id"]))
    entries.sort(key=lambda item: len(item[0]), reverse=True)
    return entries


def term_pattern(term):
    escaped = re.escape(term)
    # 英文医学短语多数由字母数字和连字符组成，用弱边界避免匹配到长词内部。
    return re.compile(rf"(?<![A-Za-z0-9_-]){escaped}(?![A-Za-z0-9_-])", re.I)


def extract_known_mentions(docs, nodes, max_mentions_per_doc=700):
    lexicon = build_lexicon(nodes)
    node_hits = defaultdict(lambda: {"count": 0, "docs": set(), "evidence": []})
    doc_to_nodes = defaultdict(set)
    for doc in docs:
        used_spans = []
        mention_count = 0
        text = doc["text"]
        for term, _lowered, node_id in lexicon:
            if mention_count >= max_mentions_per_doc:
                break
            for match in term_pattern(term).finditer(text):
                span = (match.start(), match.end())
                if any(not (span[1] <= old[0] or span[0] >= old[1]) for old in used_spans):
                    continue
                used_spans.append(span)
                node_hits[node_id]["count"] += 1
                node_hits[node_id]["docs"].add(doc["doc_id"])
                if len(node_hits[node_id]["evidence"]) < 5:
                    node_hits[node_id]["evidence"].append(snippet(text, match.start(), match.end()))
                doc_to_nodes[doc["doc_id"]].add(node_id)
                mention_count += 1
                break
    return node_hits, doc_to_nodes


def classify_candidate(label):
    lower = label.lower()
    if re.fullmatch(r"[A-Z][A-Z0-9-]{1,12}", label) or lower.endswith(" gene"):
        return "Gene/Genome"
    if any(token in lower for token in ["mutation", "mutant", "amplification", "fusion", "loss", "deletion", "variant", "methylation"]):
        return "Molecular Abnormality"
    if any(token in lower for token in ["regimen", "therapy", "chemotherapy", "immunotherapy", "treatment", "inhibitor"]):
        return "Treatment Regimen"
    if any(token in lower for token in ["stage", "tnm", "grade"]):
        return "Stage/Grade"
    if any(token in lower for token in ["colon", "rectum", "rectal", "colorectal", "liver", "lymph node", "peritoneum"]):
        return "Anatomic Site" if "carcinoma" not in lower and "cancer" not in lower else "Disease Concept"
    if any(token in lower for token in ["msi", "mmr", "microsatellite", "tmb", "hypermutated"]):
        return "Molecular Disease Subtype"
    if any(token in lower for token in ["syndrome", "disease", "polyposis"]):
        return "Associated Disease/Neoplasm"
    if any(token in lower for token in ["cell", "epithelial"]):
        return "Cell"
    if any(token in lower for token in ["mucosa", "tissue"]):
        return "Tissue"
    return "Other Entity"


def extract_rule_candidates(docs, existing_terms):
    patterns = [
        r"\b(?:MSI-H|MSI-L|MSS|dMMR|pMMR|TMB-H|ctDNA|CEA|CA19-9)\b",
        r"\b[A-Z][A-Z0-9]{1,8}(?:\s+(?:mutation|mutant|amplification|fusion|loss|deletion|variant|methylation))\b",
        r"\b(?:BRAF|KRAS|NRAS|HER2|ERBB2|PIK3CA|NTRK|ALK|ROS1)\s+[A-Z][0-9]{1,4}[A-Z]?\b",
        r"\b[A-Z][A-Za-z0-9-]+(?:\s+[A-Z][A-Za-z0-9-]+){0,4}\s+(?:regimen|therapy|inhibitor)\b",
        r"\b(?:left-sided|right-sided|metastatic|locally advanced|refractory|recurrent)\s+colorectal\s+(?:cancer|carcinoma)\b",
    ]
    compiled = [re.compile(p, re.I) for p in patterns]
    hits = defaultdict(lambda: {"count": 0, "docs": set(), "evidence": []})
    existing_norms = {normalize(term) for term in existing_terms}
    for doc in docs:
        for pattern in compiled:
            for match in pattern.finditer(doc["text"]):
                label = re.sub(r"\s+", " ", match.group(0)).strip()
                if normalize(label) in existing_norms:
                    continue
                data = hits[label]
                data["count"] += 1
                data["docs"].add(doc["doc_id"])
                if len(data["evidence"]) < 5:
                    data["evidence"].append(snippet(doc["text"], match.start(), match.end()))
    return hits


def load_config(config_path):
    if not config_path:
        return {}
    with Path(config_path).open(encoding="utf-8") as handle:
        return json.load(handle)


def call_chat_api(config, messages, temperature=0, json_mode=True, max_tokens=None):
    llm = config.get("llm", {})
    base_url = llm.get("base_url", "").rstrip("/")
    model = llm.get("model", "")
    api_key = llm.get("api_key")
    api_key_env = llm.get("api_key_env", "")
    if not api_key and api_key_env:
        # 兼容两种写法：api_key_env 可以是环境变量名，也可以是用户误填的真实 key。
        api_key = api_key_env if api_key_env.startswith("sk-") else os.getenv(api_key_env)
    if not base_url or not model or not api_key:
        raise RuntimeError("LLM API 未配置完整：需要 base_url、model、api_key 或 api_key_env。")
    endpoint = base_url
    if not endpoint.endswith("/chat/completions"):
        endpoint = endpoint + "/chat/completions"
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
    }
    if max_tokens:
        payload["max_tokens"] = max_tokens
    if json_mode:
        payload["response_format"] = {"type": "json_object"}
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        endpoint,
        data=data,
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        method="POST",
    )
    timeout = int(llm.get("timeout_seconds", 240))
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            result = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="ignore")
        if json_mode and ("response_format" in body or exc.code in {400, 422}):
            return call_chat_api(config, messages, temperature=temperature, json_mode=False, max_tokens=max_tokens)
        raise RuntimeError(f"LLM API HTTP {exc.code}: {body[:800]}") from exc
    return result["choices"][0]["message"]["content"]


def parse_json_object(text):
    text = str(text or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(text[start : end + 1])
        raise


def pubmed_text(element, path, default=""):
    found = element.find(path)
    if found is None:
        return default
    return "".join(found.itertext()).strip() or default


def search_pubmed(query, retmax=5):
    term = f"({query}) AND colorectal cancer"
    params = urllib.parse.urlencode(
        {
            "db": "pubmed",
            "term": term,
            "retmode": "json",
            "retmax": str(retmax),
            "sort": "relevance",
        }
    )
    search_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?{params}"
    try:
        with urllib.request.urlopen(search_url, timeout=30) as response:
            result = json.loads(response.read().decode("utf-8"))
        ids = result.get("esearchresult", {}).get("idlist", [])
        if not ids:
            return []
        fetch_params = urllib.parse.urlencode(
            {"db": "pubmed", "id": ",".join(ids), "retmode": "xml"}
        )
        fetch_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?{fetch_params}"
        with urllib.request.urlopen(fetch_url, timeout=45) as response:
            xml_text = response.read().decode("utf-8", errors="ignore")
    except Exception as exc:
        return [{"source": "PubMed", "error": str(exc), "query": term}]

    import xml.etree.ElementTree as ET

    root = ET.fromstring(xml_text)
    articles = []
    for article in root.findall(".//PubmedArticle"):
        pmid = pubmed_text(article, ".//PMID")
        title = pubmed_text(article, ".//ArticleTitle")
        journal = pubmed_text(article, ".//Journal/Title")
        year = pubmed_text(article, ".//PubDate/Year")
        if not year:
            year = pubmed_text(article, ".//PubDate/MedlineDate")
        abstract_parts = [
            "".join(part.itertext()).strip()
            for part in article.findall(".//Abstract/AbstractText")
            if "".join(part.itertext()).strip()
        ]
        abstract = " ".join(abstract_parts)
        doi = ""
        for aid in article.findall(".//ArticleId"):
            if aid.attrib.get("IdType") == "doi":
                doi = "".join(aid.itertext()).strip()
                break
        articles.append(
            {
                "source": "PubMed",
                "pmid": pmid,
                "title": title,
                "journal": journal,
                "year": year,
                "doi": doi,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else "",
                "abstract": clean_text(abstract)[:1600],
            }
        )
    return articles


def llm_extract(docs, config, chunk_chars=7000):
    if not config.get("llm", {}).get("enabled", False):
        return {"entities": [], "relations": [], "errors": []}
    system = (
        "你是肿瘤学知识图谱抽取器。只输出 JSON，不要解释。"
        "只抽取原文明确出现、且与结直肠癌相关的概念/实体/关系。"
        "不要扩写常识，不要编造。实体最多12个，关系最多15条。"
        "definition和evidence必须很短。"
    )
    schema = {
        "entities": [
            {
                "label": "实体名",
                "category": "Gene/Genome|Molecular Abnormality|Treatment Regimen|Disease Concept|Stage/Grade|Anatomic Site|Finding/Clinical Attribute|Associated Disease/Neoplasm|Other Entity",
                "definition": "一句话说明",
                "aliases": ["别名"],
                "evidence": "原文证据短句",
            }
        ],
        "relations": [
            {
                "source": "实体名A",
                "relation_type": "LLM_RELATED_TO|LLM_ASSOCIATED_WITH|LLM_CAUSES_OR_DRIVES|LLM_TREATS|LLM_BIOMARKER_OF|LLM_LOCATED_IN|LLM_SUBTYPE_OF",
                "target": "实体名B",
                "evidence": "原文证据短句",
            }
        ],
    }
    all_entities = []
    all_relations = []
    errors = []
    for doc in docs:
        text = doc["text"]
        chunks = prepare_llm_chunks(text, chunk_chars=chunk_chars)
        for index, chunk in enumerate(chunks):
            user = (
                f"资料名：{doc['doc_id']}\n\n"
                f"输出必须是 JSON 对象，结构：{json.dumps(schema, ensure_ascii=False)}\n\n"
                f"原文：\n{chunk}"
            )
            try:
                content = call_chat_api(
                    config,
                    [{"role": "system", "content": system}, {"role": "user", "content": user}],
                    max_tokens=6000,
                )
                parsed = parse_json_object(content)
            except Exception as exc:
                errors.append(
                    {
                        "doc_id": doc["doc_id"],
                        "chunk_index": index,
                        "error": str(exc)[:1000],
                        "raw_preview": locals().get("content", "")[:800] if "content" in locals() else "",
                    }
                )
                continue
            for entity in parsed.get("entities", []):
                entity["source_doc"] = doc["doc_id"]
                all_entities.append(entity)
            for relation in parsed.get("relations", []):
                relation["source_doc"] = doc["doc_id"]
                all_relations.append(relation)
    return {"entities": all_entities, "relations": all_relations, "errors": errors}


def prepare_llm_chunks(text, chunk_chars=7000, max_chunks=3):
    text = clean_text(text)
    if not text:
        return []
    anchors = [
        r"colorectal",
        r"\bcolon\b",
        r"\brectal\b",
        r"\bcancer\b",
        r"\bcarcinoma\b",
        r"\bCRC\b",
        r"\bMSI[- ]?H\b",
        r"\bdMMR\b",
        r"\bpMMR\b",
        r"\bMSS\b",
        r"\bKRAS\b",
        r"\bBRAF\b",
        r"\bNRAS\b",
        r"\bAPC\b",
        r"\bTP53\b",
        r"\bHER2\b",
        r"\bERBB2\b",
        r"\bPD-?1\b",
        r"\bPD-L1\b",
        r"\bCTLA-4\b",
        r"\bpembrolizumab\b",
        r"\bnivolumab\b",
        r"\bcetuximab\b",
        r"\bpanitumumab\b",
        r"\bbevacizumab\b",
        r"\bchemotherapy\b",
        r"\bimmunotherapy\b",
        r"\btargeted\b",
        r"\bmetastatic\b",
        r"\brecurrent\b",
    ]
    positions = []
    for pattern in anchors:
        for match in re.finditer(pattern, text, flags=re.I):
            positions.append(match.start())
    positions = sorted(set(positions))
    snippets = []
    used_ranges = []

    def add_span(start, end):
        start = max(0, start)
        end = min(len(text), end)
        if end <= start:
            return
        for old_start, old_end in used_ranges:
            if not (end <= old_start or start >= old_end):
                return
        used_ranges.append((start, end))
        snippets.append(text[start:end].strip())

    add_span(0, min(len(text), chunk_chars))
    for pos in positions[:20]:
        add_span(pos - chunk_chars // 3, pos + chunk_chars)
        if len(snippets) >= max_chunks:
            break
    if len(snippets) < max_chunks and len(text) > chunk_chars:
        step = max(chunk_chars, len(text) // max_chunks)
        for start in range(0, len(text), step):
            add_span(start, start + chunk_chars)
            if len(snippets) >= max_chunks:
                break
    combined = []
    total = 0
    for idx, snippet in enumerate(snippets[:max_chunks], start=1):
        block = f"### 片段 {idx}\n{snippet}"
        if total + len(block) > chunk_chars * max_chunks:
            break
        combined.append(block)
        total += len(block)
    return combined


def llm_summarize_documents(docs, config, max_chars=12000):
    if not config.get("llm", {}).get("enabled", False):
        return {"enabled": False, "brief": "", "analysis": "", "error": ""}
    joined = []
    for doc in docs[:8]:
        text = doc.get("text", "")
        snippets = prepare_llm_chunks(text, chunk_chars=max_chars // max(1, min(len(docs), 8)), max_chunks=2)
        joined.append(f"## {doc.get('title') or doc.get('doc_id')}\n来源：{doc.get('path', '')}\n" + "\n\n".join(snippets)[: max_chars // max(1, min(len(docs), 8))])
    system = (
        "你是结直肠癌知识图谱资料分析助手。只根据用户给定资料总结，不要编造。"
        "输出 JSON，字段为 brief、analysis。brief 用 3-5 条短句给学生看；"
        "analysis 用 Markdown，包含资料类型/核心问题/关键概念实体/关系线索/可并入知识图谱的信息/需要人工核对处。"
    )
    user = "请总结并分析以下新资料，用于结直肠癌知识图谱更新：\n\n" + "\n\n".join(joined)
    try:
        content = call_chat_api(
            config,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            temperature=0,
            max_tokens=5000,
        )
        parsed = parse_json_object(content)
        return {
            "enabled": True,
            "brief": clean_text(parsed.get("brief", "")),
            "analysis": str(parsed.get("analysis", "")).strip(),
            "error": "",
        }
    except Exception as exc:
        return {"enabled": True, "brief": "", "analysis": "", "error": str(exc)[:1000]}


def make_node(base=None, **updates):
    row = {field: "" for field in NODE_FIELDS}
    if base:
        row.update({field: base.get(field, "") for field in NODE_FIELDS})
    row.update(updates)
    if not row.get("research_axis"):
        row["research_axis"] = CATEGORY_TO_AXIS.get(row.get("category", ""), "Other")
    return row


def make_edge(source, source_label, relation_type, target, target_label, docs="", evidence=""):
    return {
        "source": source,
        "source_label": source_label,
        "relation_code": "",
        "relation_type": relation_type,
        "relation_label_cn": RELATION_LABELS.get(relation_type, relation_type),
        "target": target,
        "target_label": target_label,
        "relation_group": "extracted",
        "direction": "out",
        "research_tier": "incremental",
        "source_docs": docs,
        "evidence": evidence,
    }


def build_mini_kg(
    input_path,
    kg_dir,
    output_dir,
    config=None,
    use_llm=False,
    obsidian_vault=None,
    obsidian_folder="结直肠癌知识图谱",
    obsidian_open=False,
    summarize=False,
):
    config = config or {}
    extract_config = copy.deepcopy(config)
    summary_config = copy.deepcopy(config)
    extract_config.setdefault("llm", {})["enabled"] = bool(use_llm)
    summary_config.setdefault("llm", {})["enabled"] = True
    reader_config = extract_config if use_llm else summary_config
    docs, parse_errors = load_documents(input_path, config=reader_config, use_llm=bool(use_llm or summarize), return_errors=True)
    if not docs:
        detail = format_parse_errors(parse_errors)
        hint = (
            "未读取到可处理资料。\n"
            f"支持格式：{', '.join(sorted(SUPPORTED_EXTS))}\n"
            "PDF：文字版 PDF 可直接读取；扫描版 PDF 需要 OCR。\n"
            "图片：需要本机 Tesseract OCR，或勾选“大模型增强”后使用支持图片输入的模型读取。"
        )
        raise RuntimeError(hint + (f"\n\n文件级诊断：\n{detail}" if detail else ""))
    base_nodes, base_edges = load_base_kg(kg_dir)
    known_hits, doc_to_nodes = extract_known_mentions(docs, base_nodes)
    existing_terms = []
    for node in base_nodes.values():
        existing_terms.append(node.get("label", ""))
        existing_terms.extend(node.get("aliases", "").split("|"))
    candidate_hits = extract_rule_candidates(docs, existing_terms)
    llm_result = llm_extract(docs, extract_config) if extract_config.get("llm", {}).get("enabled", False) else {"entities": [], "relations": []}
    summary_result = llm_summarize_documents(docs, summary_config) if summarize else {"enabled": False, "brief": "", "analysis": "", "error": ""}

    mini_nodes = {}
    for node_id, hit in known_hits.items():
        base = base_nodes[node_id]
        mini_nodes[node_id] = make_node(
            base,
            kg_status="existing",
            source_docs="|".join(sorted(hit["docs"])),
            evidence=" || ".join(hit["evidence"]),
        )

    label_to_id = {normalize(row["label"]): row["id"] for row in mini_nodes.values()}
    base_term_to_id = {}
    for base_id, base in base_nodes.items():
        for term in [base.get("label", "")] + base.get("aliases", "").split("|"):
            term_norm = normalize(term)
            if term_norm:
                base_term_to_id.setdefault(term_norm, base_id)
    for label, hit in candidate_hits.items():
        node_id = stable_new_id(label)
        category = classify_candidate(label)
        mini_nodes[node_id] = make_node(
            id=node_id,
            label=label,
            node_kind="Entity",
            in_colorectal_tree="N",
            category=category,
            research_axis=CATEGORY_TO_AXIS.get(category, "Other"),
            semantic_types="",
            definition="规则抽取的新候选实体，需人工确认。",
            aliases=label,
            alias_count="1",
            kg_status="new_candidate_rule",
            source_docs="|".join(sorted(hit["docs"])),
            evidence=" || ".join(hit["evidence"]),
        )
        label_to_id[normalize(label)] = node_id

    for entity in llm_result.get("entities", []):
        label = str(entity.get("label", "")).strip()
        if not label:
            continue
        norm = normalize(label)
        node_id = label_to_id.get(norm) or base_term_to_id.get(norm) or stable_new_id(label)
        category = entity.get("category") or classify_candidate(label)
        aliases = [label] + [str(alias) for alias in entity.get("aliases", []) if alias]
        if node_id in base_nodes:
            if node_id not in mini_nodes:
                mini_nodes[node_id] = make_node(base_nodes[node_id])
            old_existing = mini_nodes[node_id]
            old_existing["kg_status"] = "existing"
            old_existing["source_docs"] = "|".join(
                sorted(
                    set(
                        filter(
                            None,
                            old_existing.get("source_docs", "").split("|")
                            + [entity.get("source_doc", "")],
                        )
                    )
                )
            )
            old_existing["evidence"] = " || ".join(
                part for part in [old_existing.get("evidence", ""), entity.get("evidence", "")] if part
            )
            label_to_id[norm] = node_id
            continue
        old = mini_nodes.get(node_id, {})
        docs_joined = "|".join(sorted(set((old.get("source_docs", "").split("|") if old.get("source_docs") else []) + [entity.get("source_doc", "")])))
        evidence = " || ".join([part for part in [old.get("evidence", ""), entity.get("evidence", "")] if part])
        mini_nodes[node_id] = make_node(
            old,
            id=node_id,
            label=old.get("label") or label,
            node_kind="Entity",
            in_colorectal_tree="N",
            category=old.get("category") or category,
            research_axis=CATEGORY_TO_AXIS.get(old.get("category") or category, "Other"),
            definition=old.get("definition") if old.get("definition") and "规则抽取" not in old.get("definition", "") else entity.get("definition", "模型抽取的新候选实体，需人工确认。"),
            aliases="|".join(sorted(set((old.get("aliases", "").split("|") if old.get("aliases") else []) + aliases))),
            alias_count=str(len(set((old.get("aliases", "").split("|") if old.get("aliases") else []) + aliases))),
            kg_status="new_candidate_llm" if old.get("kg_status") != "new_candidate_rule" else "new_candidate_rule_llm",
            source_docs=docs_joined,
            evidence=evidence,
        )
        label_to_id[norm] = node_id

    mini_edges = []
    base_edge_keys = {(edge["source"], edge["relation_type"], edge["target"]) for edge in base_edges}
    mini_ids = set(mini_nodes)
    for edge in base_edges:
        if edge["source"] in mini_ids and edge["target"] in mini_ids:
            row = {field: edge.get(field, "") for field in EDGE_FIELDS}
            row["source_docs"] = ""
            row["evidence"] = ""
            mini_edges.append(row)

    for doc in docs:
        ids = sorted(doc_to_nodes.get(doc["doc_id"], set()) & mini_ids, key=lambda nid: mini_nodes[nid]["label"])
        ids = ids[:80]
        for idx, source in enumerate(ids):
            for target in ids[idx + 1 : idx + 9]:
                if (source, "MENTIONED_WITH", target) in base_edge_keys:
                    continue
                mini_edges.append(
                    make_edge(
                        source,
                        mini_nodes[source]["label"],
                        "MENTIONED_WITH",
                        target,
                        mini_nodes[target]["label"],
                        doc["doc_id"],
                        "同一资料内共现；表示线索，不表示因果。",
                    )
                )

    for relation in llm_result.get("relations", []):
        s_label = str(relation.get("source", "")).strip()
        t_label = str(relation.get("target", "")).strip()
        if not s_label or not t_label:
            continue
        source = label_to_id.get(normalize(s_label))
        target = label_to_id.get(normalize(t_label))
        if not source or not target:
            continue
        relation_type = relation.get("relation_type") or "LLM_RELATED_TO"
        mini_edges.append(
            make_edge(
                source,
                mini_nodes[source]["label"],
                relation_type,
                target,
                mini_nodes[target]["label"],
                relation.get("source_doc", ""),
                relation.get("evidence", ""),
            )
        )

    deduped_edges = []
    seen = set()
    for edge in mini_edges:
        key = (edge["source"], edge["relation_type"], edge["target"], edge.get("source_docs", ""))
        if key in seen:
            continue
        seen.add(key)
        deduped_edges.append(edge)

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    write_csv(output_dir / "mini_nodes.csv", sorted(mini_nodes.values(), key=lambda n: (n["kg_status"], n["label"])), NODE_FIELDS)
    write_csv(output_dir / "mini_edges.csv", sorted(deduped_edges, key=lambda e: (e["source_label"], e["relation_type"], e["target_label"])), EDGE_FIELDS)
    write_graphml(mini_nodes, deduped_edges, output_dir / "mini_graph.graphml")
    write_cytoscape(mini_nodes, deduped_edges, output_dir / "mini_cytoscape.json")
    write_mini_html(mini_nodes, deduped_edges, output_dir / "mini_browser.html", "新增资料小图谱")
    write_tree_html(mini_nodes, deduped_edges, output_dir / "mini_tree.html", "新增资料概念-实体树")

    report = {
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "input_path": str(Path(input_path).resolve()),
        "document_count": len(docs),
        "known_node_count": sum(1 for n in mini_nodes.values() if n["kg_status"] == "existing"),
        "new_candidate_count": sum(1 for n in mini_nodes.values() if n["kg_status"] != "existing"),
        "edge_count": len(deduped_edges),
        "llm_enabled": bool(extract_config.get("llm", {}).get("enabled", False)),
        "llm_entity_count": len(llm_result.get("entities", [])),
        "llm_relation_count": len(llm_result.get("relations", [])),
        "llm_error_count": len(llm_result.get("errors", [])),
        "llm_errors": llm_result.get("errors", []),
        "parse_error_count": len(parse_errors),
        "parse_errors": parse_errors,
        "summary_enabled": bool(summarize),
        "summary_brief": summary_result.get("brief", ""),
        "summary_analysis": summary_result.get("analysis", ""),
        "summary_error": summary_result.get("error", ""),
    }
    report["title"] = guess_title(" ".join(doc.get("title", "") for doc in docs), "新资料小图谱")
    report["parse_notes"] = [
        f"{doc['doc_id']}: {doc['file_type']}" + (f" 解析备注：{doc['parse_error']}" if doc.get("parse_error") else "")
        for doc in docs
        if doc.get("parse_error")
    ]
    report["parse_notes"].extend(format_parse_errors(parse_errors).splitlines())
    (output_dir / "mini_report.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    write_tree_json(mini_nodes, deduped_edges, output_dir / "mini_tree.json")
    markdown_text = build_markdown_report(report, docs, list(mini_nodes.values()), deduped_edges)
    write_markdown(output_dir / "mini_note.md", markdown_text)
    if obsidian_vault:
        obsidian_result = write_obsidian_note(
            obsidian_vault,
            report["title"] or f"mini_{time.strftime('%Y%m%d_%H%M%S')}",
            markdown_text,
            folder=obsidian_folder,
            open_after=obsidian_open,
            extra_links={
                "小图谱 HTML": str((output_dir / "mini_browser.html").resolve()),
                "树状结构 HTML": str((output_dir / "mini_tree.html").resolve()),
                "小图谱数据": str((output_dir / "mini_nodes.csv").resolve()),
            },
        )
        report["obsidian_note_path"] = obsidian_result.get("note_path", "")
        report["obsidian_uri"] = obsidian_result.get("uri", "")
    return report


def merge_kg(
    base_kg_dir,
    mini_dir,
    output_dir,
    accept_candidates=False,
    obsidian_vault=None,
    obsidian_folder="结直肠癌知识图谱",
    obsidian_open=False,
):
    base_kg_dir = Path(base_kg_dir)
    mini_dir = Path(mini_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_dir = output_dir / f"_backup_source_{timestamp}"
    backup_dir.mkdir(parents=True, exist_ok=True)
    for file in base_kg_dir.glob("*"):
        if file.is_file():
            shutil.copy2(file, backup_dir / file.name)

    base_nodes, base_edges = load_base_kg(base_kg_dir)
    mini_nodes = {row["id"]: row for row in read_csv(mini_dir / "mini_nodes.csv")}
    mini_edges = read_csv(mini_dir / "mini_edges.csv")

    merged_nodes = {nid: dict(row) for nid, row in base_nodes.items()}
    added_nodes = 0
    touched_nodes = 0
    for node_id, node in mini_nodes.items():
        if node_id in merged_nodes:
            old = merged_nodes[node_id]
            old["kg_status"] = "existing_updated"
            old["source_docs"] = "|".join(sorted(set(filter(None, old.get("source_docs", "").split("|") + node.get("source_docs", "").split("|")))))
            old["evidence"] = " || ".join(filter(None, [old.get("evidence", ""), node.get("evidence", "")]))[:1800]
            touched_nodes += 1
        elif accept_candidates:
            node["kg_status"] = node.get("kg_status") or "new_candidate_accepted"
            merged_nodes[node_id] = node
            added_nodes += 1

    merged_edges = []
    seen = set()
    for edge in base_edges + mini_edges:
        if edge["source"] not in merged_nodes or edge["target"] not in merged_nodes:
            continue
        key = (edge["source"], edge["relation_type"], edge["target"])
        if key in seen:
            continue
        seen.add(key)
        merged_edges.append({field: edge.get(field, "") for field in EDGE_FIELDS})

    write_csv(output_dir / "kg_nodes.csv", sorted(merged_nodes.values(), key=lambda n: (n.get("node_kind", ""), n.get("category", ""), n.get("label", ""))), NODE_FIELDS)
    write_csv(output_dir / "kg_edges_research_core.csv", sorted(merged_edges, key=lambda e: (e["source_label"], e["relation_type"], e["target_label"])), EDGE_FIELDS)
    write_csv(output_dir / "kg_edges.csv", sorted(merged_edges, key=lambda e: (e["source_label"], e["relation_type"], e["target_label"])), EDGE_FIELDS)
    write_graphml(merged_nodes, merged_edges, output_dir / "kg_graph.graphml")
    write_cytoscape(merged_nodes, merged_edges, output_dir / "kg_cytoscape.json")
    write_main_style_html(merged_nodes, merged_edges, output_dir / "kg_browser.html")
    write_tree_json(merged_nodes, merged_edges, output_dir / "kg_tree.json")
    write_tree_html(merged_nodes, merged_edges, output_dir / "kg_tree.html", "更新后的结直肠癌概念-实体树")
    summary = {
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "base_kg_dir": str(base_kg_dir.resolve()),
        "mini_dir": str(mini_dir.resolve()),
        "accept_candidates": accept_candidates,
        "node_count": len(merged_nodes),
        "edge_count": len(merged_edges),
        "added_nodes": added_nodes,
        "touched_existing_nodes": touched_nodes,
        "backup_dir": str(backup_dir.resolve()),
    }
    write_markdown(
        output_dir / "kg_note.md",
        build_markdown_report(
            {
                "title": "更新后的结直肠癌知识图谱",
                "generated_at": summary["generated_at"],
                "input_path": str(mini_dir.resolve()),
                "document_count": 0,
                "known_node_count": touched_nodes,
                "new_candidate_count": added_nodes,
                "edge_count": len(merged_edges),
                "parse_notes": [],
            },
            [],
            list(merged_nodes.values()),
            merged_edges,
        ),
    )
    if obsidian_vault:
        obsidian_result = write_obsidian_note(
            obsidian_vault,
            "更新后的结直肠癌知识图谱",
            (output_dir / "kg_note.md").read_text(encoding="utf-8"),
            folder=obsidian_folder,
            open_after=obsidian_open,
            extra_links={
                "更新图谱 HTML": str((output_dir / "kg_browser.html").resolve()),
                "树状结构 HTML": str((output_dir / "kg_tree.html").resolve()),
                "图谱数据": str((output_dir / "kg_nodes.csv").resolve()),
            },
        )
        summary["obsidian_note_path"] = obsidian_result.get("note_path", "")
        summary["obsidian_uri"] = obsidian_result.get("uri", "")
    (output_dir / "merge_report.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def retrieve_context(question, kg_dir, top_k=20):
    nodes, edges = load_base_kg(kg_dir)
    q_terms = [term for term in re.split(r"[^A-Za-z0-9\u4e00-\u9fff-]+", normalize(question)) if len(term) >= 2]
    generic_terms = {
        "cancer",
        "carcinoma",
        "colorectal",
        "colon",
        "rectal",
        "tumor",
        "neoplasm",
        "什么",
        "关系",
        "有什么",
    }
    rare_terms = [term for term in q_terms if term not in generic_terms]
    scored = []
    for node in nodes.values():
        label = normalize(node.get("label", ""))
        aliases = normalize(node.get("aliases", ""))
        definition = normalize(node.get("definition", ""))
        category = normalize(node.get("category", ""))
        blob = " ".join([label, aliases, definition, category])
        score = 0
        for term in q_terms:
            if term not in blob:
                continue
            weight = 1 if term in generic_terms else 4
            if term in label:
                score += 7 * weight
            elif term in aliases:
                score += 4 * weight
            elif term in category:
                score += 2 * weight
            else:
                score += weight
        if rare_terms and not any(term in label or term in aliases for term in rare_terms):
            score *= 0.35
        if score:
            scored.append((score + min(5, int(node.get("alias_count") or 0) * 0.1), node))
    scored.sort(key=lambda item: item[0], reverse=True)
    selected = [node for _score, node in scored[:top_k]]
    selected_ids = {node["id"] for node in selected}
    related_edges = [edge for edge in edges if edge["source"] in selected_ids or edge["target"] in selected_ids][:top_k * 3]
    context = {
        "nodes": [
            {
                "id": node["id"],
                "label": node["label"],
                "kind": node.get("node_kind", ""),
                "category": node.get("category", ""),
                "definition": node.get("definition", ""),
                "aliases": node.get("aliases", ""),
            }
            for node in selected
        ],
        "edges": [
            {
                "source": edge["source_label"],
                "relation": edge["relation_label_cn"] or edge["relation_type"],
                "target": edge["target_label"],
            }
            for edge in related_edges
        ],
    }
    return context


def normalize_chat_history(chat_history, max_turns=6):
    normalized = []
    for item in chat_history or []:
        if not isinstance(item, dict):
            continue
        role = str(item.get("role", "")).strip().lower()
        if role not in {"user", "assistant"}:
            continue
        content = clean_text(item.get("content", ""))
        if not content:
            continue
        normalized.append({"role": role, "content": content[:1800]})
    return normalized[-max(0, int(max_turns)) * 2 :]


def format_chat_history(chat_history):
    if not chat_history:
        return "无"
    lines = []
    for item in chat_history:
        role = "用户" if item["role"] == "user" else "助手"
        lines.append(f"{role}：{item['content']}")
    return "\n".join(lines)


def answer_question(
    question,
    kg_dir,
    config=None,
    top_k=20,
    use_external=False,
    external_k=5,
    answer_mode="graph",
    chat_history=None,
    memory_enabled=True,
    memory_turns=6,
):
    config = config or {}
    history = normalize_chat_history(chat_history, memory_turns) if memory_enabled else []
    history_user_text = " ".join(item["content"] for item in history if item["role"] == "user")
    retrieval_question = f"{history_user_text} {question}".strip() if history_user_text else question
    context = retrieve_context(retrieval_question, kg_dir, top_k)
    if answer_mode in {"literature", "comprehensive"}:
        use_external = True
    external_context = search_pubmed(question, external_k) if use_external else []
    if not config.get("llm", {}).get("enabled", False):
        return {
            "mode": "retrieval_only",
            "answer": "未启用第三方大模型 API。以下是从图谱检索到的候选上下文，可用于人工判断或接入模型后生成自然语言答案。",
            "context": context,
            "external_context": external_context,
            "answer_mode": answer_mode,
            "chat_history_used": len(history),
        }
    if answer_mode == "literature":
        system = (
            "你是结直肠癌文献综述助手。必须优先总结给定 PubMed 文献上下文，并结合本地图谱做术语和关系校准。"
            "不要把模型常识冒充为文献证据。"
            "请输出 JSON，字段包括 answer、pubmed_summary、graph_basis、third_party_basis、conflicts_or_uncertainty、limitations。"
            "pubmed_summary 按 PMID 分条总结关键发现；third_party_basis 必须列 PMID/URL；graph_basis 只写来自图谱的依据。"
        )
    elif answer_mode == "comprehensive":
        system = (
            "你是结直肠癌综合问答助手，可以像普通医学 AI 一样结合自身已学知识、给定本地图谱、给定 PubMed 文献进行回答。"
            "必须区分信息来源：本地图谱、PubMed第三方资料、AI综合医学知识。"
            "不要声称 AI综合医学知识来自实时网络；实时/第三方证据只能来自提供的 PubMed 上下文。"
            "请输出 JSON，字段包括 answer、graph_basis、pubmed_summary、third_party_basis、ai_general_knowledge、conflicts_or_uncertainty、limitations。"
        )
    else:
        system = (
            "你是结直肠癌知识图谱问答助手。优先基于给定图谱上下文回答。"
            "如果提供了第三方资料上下文，可以使用，但必须明确标注哪些信息来自第三方资料。"
            "请输出 JSON，字段包括 answer、graph_basis、third_party_basis、limitations。"
            "graph_basis 写图谱依据；third_party_basis 写第三方资料依据及 PMID/URL；limitations 写证据不足或需人工核查处。"
        )
    user = (
        f"问题：{question}\n\n"
        f"回答模式：{answer_mode}\n\n"
        f"本轮之前的对话记忆：\n{format_chat_history(history)}\n\n"
        f"图谱上下文：\n{json.dumps(context, ensure_ascii=False)}\n\n"
        f"第三方资料上下文：\n{json.dumps(external_context, ensure_ascii=False)}"
    )
    content = call_chat_api(config, [{"role": "system", "content": system}, {"role": "user", "content": user}], temperature=0)
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError:
        parsed = {"answer": content}
    parsed["context"] = context
    parsed["external_context"] = external_context
    parsed["mode"] = "llm_answer"
    parsed["answer_mode"] = answer_mode
    parsed["chat_history_used"] = len(history)
    return parsed


def write_graphml(nodes, edges, path):
    node_keys = ["label", "node_kind", "category", "research_axis", "kg_status", "definition", "aliases", "source_docs"]
    edge_keys = ["relation_type", "relation_label_cn", "research_tier", "source_docs", "evidence"]
    lines = ['<?xml version="1.0" encoding="UTF-8"?>', '<graphml xmlns="http://graphml.graphdrawing.org/xmlns">']
    for key in node_keys:
        lines.append(f'<key id="n_{key}" for="node" attr.name="{key}" attr.type="string"/>')
    for key in edge_keys:
        lines.append(f'<key id="e_{key}" for="edge" attr.name="{key}" attr.type="string"/>')
    lines.append('<graph id="IncrementalKG" edgedefault="directed">')
    for node in nodes.values():
        lines.append(f'<node id="{xml_escape(node["id"])}">')
        for key in node_keys:
            lines.append(f'<data key="n_{key}">{xml_escape(str(node.get(key, "")))}</data>')
        lines.append("</node>")
    for idx, edge in enumerate(edges):
        lines.append(f'<edge id="e{idx}" source="{xml_escape(edge["source"])}" target="{xml_escape(edge["target"])}">')
        for key in edge_keys:
            lines.append(f'<data key="e_{key}">{xml_escape(str(edge.get(key, "")))}</data>')
        lines.append("</edge>")
    lines.extend(["</graph>", "</graphml>"])
    Path(path).write_text("\n".join(lines), encoding="utf-8")


def write_cytoscape(nodes, edges, path):
    elements = {
        "nodes": [{"data": node} for node in nodes.values()],
        "edges": [{"data": {"id": f"{e['source']}__{e['relation_type']}__{e['target']}", **e}} for e in edges],
    }
    Path(path).write_text(json.dumps(elements, ensure_ascii=False, indent=2), encoding="utf-8")


def markdown_escape(text):
    return clean_text(text).replace("|", "\\|")


def build_markdown_report(report, docs, nodes, edges):
    lines = []
    lines.append(f"# {report.get('title') or '新资料抽取笔记'}")
    lines.append("")
    lines.append(f"- 生成时间：{report.get('generated_at', '')}")
    lines.append(f"- 输入来源：{report.get('input_path', '')}")
    lines.append(f"- 文档数量：{report.get('document_count', 0)}")
    lines.append(f"- 已知节点：{report.get('known_node_count', 0)}")
    lines.append(f"- 新增候选：{report.get('new_candidate_count', 0)}")
    lines.append(f"- 关系数：{report.get('edge_count', 0)}")
    if report.get("summary_brief") or report.get("summary_analysis") or report.get("summary_error"):
        lines.append("")
        lines.append("## 大模型资料总结")
        if report.get("summary_brief"):
            lines.append("")
            lines.append("### 简要总结")
            for part in re.split(r"(?:\n|；|;)+", str(report.get("summary_brief", ""))):
                part = part.strip(" -•\t")
                if part:
                    lines.append(f"- {markdown_escape(part)}")
        if report.get("summary_analysis"):
            lines.append("")
            lines.append("### 分析与整合建议")
            lines.append(str(report.get("summary_analysis", "")).strip())
        if report.get("summary_error"):
            lines.append("")
            lines.append(f"- 总结失败：{markdown_escape(report.get('summary_error'))}")
    if report.get("parse_notes"):
        lines.append("")
        lines.append("## 解析备注")
        for note in report["parse_notes"]:
            lines.append(f"- {note}")
    if docs:
        lines.append("")
        lines.append("## 文档")
        for doc in docs:
            lines.append(f"- **{markdown_escape(doc.get('title') or doc['doc_id'])}**  ")
            lines.append(f"  - 类型：{doc.get('file_type', '')}")
            lines.append(f"  - 路径：{doc.get('path', '')}")
            if doc.get("parse_error"):
                lines.append(f"  - 解析备注：{doc.get('parse_error')}")
    if nodes:
        lines.append("")
        lines.append("## 节点")
        for node in nodes:
            lines.append(f"- **{markdown_escape(node.get('label', ''))}** `{node.get('id', '')}`")
            meta = []
            if node.get("category"):
                meta.append(f"类别：{node.get('category')}")
            if node.get("kg_status"):
                meta.append(f"状态：{node.get('kg_status')}")
            if node.get("definition"):
                meta.append(f"定义：{markdown_escape(node.get('definition'))[:160]}")
            if meta:
                lines.append("  - " + "；".join(meta))
    if edges:
        lines.append("")
        lines.append("## 关系")
        for edge in edges[:200]:
            lines.append(
                f"- {markdown_escape(edge.get('source_label', ''))} → {markdown_escape(edge.get('relation_label_cn') or edge.get('relation_type', ''))} → {markdown_escape(edge.get('target_label', ''))}"
            )
    return "\n".join(lines).strip() + "\n"


def write_markdown(path, text):
    Path(path).write_text(text, encoding="utf-8")


def make_obsidian_uri(vault_path, note_path):
    vault_name = Path(vault_path).name
    rel = Path(note_path).relative_to(Path(vault_path)).as_posix()
    return f"obsidian://open?vault={urllib.parse.quote(vault_name)}&file={urllib.parse.quote(rel)}"


def write_obsidian_note(vault_path, note_name, markdown_text, folder="结直肠癌知识图谱", open_after=False, extra_links=None):
    vault_path = Path(vault_path)
    vault_path.mkdir(parents=True, exist_ok=True)
    target_dir = vault_path / folder
    target_dir.mkdir(parents=True, exist_ok=True)
    note_path = target_dir / f"{note_name}.md"
    links = extra_links or {}
    header = []
    if links:
        header.append("## 快速链接")
        for label, rel_path in links.items():
            header.append(f"- [{label}]({Path(rel_path).as_posix()})")
        header.append("")
    note_path.write_text("\n".join(header) + markdown_text, encoding="utf-8")
    result = {"note_path": str(note_path.resolve())}
    if open_after:
        result["uri"] = make_obsidian_uri(vault_path, note_path)
    return result


def write_tree_json(nodes, edges, path):
    Path(path).write_text(json.dumps(build_tree_payload(nodes, edges), ensure_ascii=False, indent=2), encoding="utf-8")


def tree_sort_key(node):
    depth = node.get("depth_from_root", "")
    try:
        depth_value = int(depth)
    except (TypeError, ValueError):
        depth_value = 9999
    return (depth_value, node.get("category", ""), node.get("label", ""), node.get("id", ""))


def build_tree_payload(nodes, edges):
    node_index = {node["id"]: node for node in nodes.values()}
    concept_ids = {node_id for node_id, node in node_index.items() if node.get("node_kind") == "Concept"}
    entity_ids = set(node_index) - concept_ids
    rows = []
    row_ids = set()
    concept_children = defaultdict(list)
    concept_parent = {}
    for edge in edges:
        if edge.get("relation_type") != "is_parent_of":
            continue
        source = edge.get("source")
        target = edge.get("target")
        if source in concept_ids and target in concept_ids and target not in concept_parent:
            concept_children[source].append(target)
            concept_parent[target] = source

    def add_row(row):
        row["row_id"] = row.get("row_id") or row.get("id")
        if row["row_id"] in row_ids:
            return
        row_ids.add(row["row_id"])
        rows.append(row)

    def node_row(node_id, parent="", depth=0, relation=""):
        node = node_index.get(node_id)
        if not node:
            return {}
        return {
            "row_id": node_id,
            "id": node["id"],
            "label": node["label"],
            "kind": node.get("node_kind", ""),
            "category": node.get("category", ""),
            "definition": node.get("definition", ""),
            "aliases": node.get("aliases", ""),
            "status": node.get("kg_status", ""),
            "source_docs": node.get("source_docs", ""),
            "evidence": node.get("evidence", ""),
            "parent": parent,
            "depth": depth,
            "relation": relation,
        }

    visited = set()

    def add_concept_branch(node_id, parent="", depth=0, relation="下位概念"):
        if node_id in visited or node_id not in concept_ids:
            return
        visited.add(node_id)
        add_row(node_row(node_id, parent=parent, depth=depth, relation=relation))
        for child_id in sorted(set(concept_children.get(node_id, [])), key=lambda cid: tree_sort_key(node_index.get(cid, {}))):
            add_concept_branch(child_id, parent=node_id, depth=depth + 1, relation="下位概念")

    if "C2955" in concept_ids:
        root_ids = ["C2955"] + [
            node_id
            for node_id in sorted(concept_ids, key=lambda cid: tree_sort_key(node_index[cid]))
            if node_id not in concept_parent and node_id != "C2955"
        ]
    else:
        root_ids = [node_id for node_id in sorted(concept_ids, key=lambda cid: tree_sort_key(node_index[cid])) if node_id not in concept_parent]
    if not root_ids:
        root_ids = sorted(concept_ids, key=lambda cid: tree_sort_key(node_index[cid]))[:1]

    for root_id in root_ids:
        add_concept_branch(root_id, parent="", depth=0, relation="根概念")
    for concept_id in sorted(concept_ids - visited, key=lambda cid: tree_sort_key(node_index[cid])):
        add_concept_branch(concept_id, parent="", depth=0, relation="独立概念")

    if entity_ids:
        entity_root_id = "__ENTITY_FACETS__"
        add_row(
            {
                "row_id": entity_root_id,
                "id": entity_root_id,
                "label": "实体分面（按类别）",
                "kind": "Group",
                "category": "Entity Facets",
                "definition": "把基因、分子异常、解剖部位、治疗方案、临床发现等实体按类别集中展示；右侧详情可查看实体与概念的关系。",
                "aliases": "",
                "status": "",
                "source_docs": "",
                "evidence": "",
                "parent": "",
                "depth": 0,
                "relation": "实体导航",
            }
        )
        grouped_entities = defaultdict(list)
        for entity_id in entity_ids:
            grouped_entities[node_index[entity_id].get("category", "Other Entity") or "Other Entity"].append(entity_id)
        for category in sorted(grouped_entities):
            group_id = f"__ENTITY_CAT__{hashlib.sha1(category.encode('utf-8')).hexdigest()[:10]}"
            add_row(
                {
                    "row_id": group_id,
                    "id": group_id,
                    "label": category,
                    "kind": "Group",
                    "category": "Entity Category",
                    "definition": f"{category} 类实体。",
                    "aliases": "",
                    "status": "",
                    "source_docs": "",
                    "evidence": "",
                    "parent": entity_root_id,
                    "depth": 1,
                    "relation": "实体类别",
                }
            )
            for entity_id in sorted(grouped_entities[category], key=lambda cid: tree_sort_key(node_index[cid])):
                add_row(node_row(entity_id, parent=group_id, depth=2, relation="类别成员"))

    edge_rows = []
    for idx, edge in enumerate(edges):
        source = edge.get("source", "")
        target = edge.get("target", "")
        if source not in node_index or target not in node_index:
            continue
        edge_rows.append(
            {
                "id": f"e{idx}",
                "source": source,
                "source_label": edge.get("source_label") or node_index[source].get("label", source),
                "relation": edge.get("relation_label_cn") or edge.get("relation_type", ""),
                "relation_type": edge.get("relation_type", ""),
                "target": target,
                "target_label": edge.get("target_label") or node_index[target].get("label", target),
                "source_docs": edge.get("source_docs", ""),
                "evidence": edge.get("evidence", ""),
            }
        )

    return {
        "title": "结直肠癌概念-实体树",
        "root_count": sum(1 for row in rows if not row.get("parent")),
        "concept_count": len(concept_ids),
        "entity_count": len(entity_ids),
        "row_count": len(rows),
        "edge_count": len(edge_rows),
        "rows": rows,
        "edges": edge_rows,
    }


def write_tree_html(nodes, edges, path, title="结直肠癌概念-实体树"):
    payload = build_tree_payload(nodes, edges)
    payload["title"] = title
    template = """<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>%%TITLE%%</title>
<style>
:root{--ink:#172033;--muted:#667085;--line:#d9e2ec;--bg:#f4f8fc;--blue:#2563eb;--red:#e11d48;--green:#0f766e;--orange:#f97316;--violet:#7c3aed;--slate:#475569}
*{box-sizing:border-box}
body{margin:0;font-family:Arial,"Microsoft YaHei",sans-serif;color:var(--ink);background:#f4f8fc}
header{min-height:76px;display:flex;align-items:center;justify-content:space-between;gap:16px;padding:12px 18px;background:rgba(255,255,255,.96);border-bottom:1px solid var(--line);backdrop-filter:blur(8px)}
h1{margin:0;font-size:20px}.sub{font-size:12px;color:var(--muted);margin-top:4px;line-height:1.5}
main{display:grid;grid-template-columns:260px minmax(620px,1fr) 360px;height:calc(100vh - 76px);min-height:620px}
aside{overflow:auto;background:#fff;border-right:1px solid var(--line);padding:12px}#details{border-left:1px solid var(--line);border-right:0}.center{overflow:auto;padding:12px;background:#f7fafc}
.card{background:#fff;border:1px solid var(--line);border-radius:8px;padding:12px;margin-bottom:10px}.muted{color:var(--muted);font-size:12px;line-height:1.55}
.stats{display:grid;grid-template-columns:1fr 1fr;gap:8px}.stat{border:1px solid var(--line);border-radius:8px;padding:9px;background:#f8fafc}.stat b{display:block;font-size:20px;margin-bottom:1px}
input,button{border:1px solid #cbd5e1;border-radius:8px;background:#fff;padding:8px 10px;font-size:13px}button{cursor:pointer}.toolbar{display:flex;gap:8px;align-items:center;flex-wrap:wrap}
.legend{display:flex;gap:6px;flex-wrap:wrap;margin-top:8px}.legend span,.pill{display:inline-flex;align-items:center;gap:5px;margin:3px 4px 3px 0;padding:3px 8px;background:#eef2f7;border-radius:999px;font-size:11px}.dot{width:8px;height:8px;border-radius:50%;display:inline-block}
.tree-wrap{background:#fff;border:1px solid var(--line);border-radius:8px;overflow:hidden;box-shadow:0 8px 24px rgba(15,23,42,.04)}.tree-head{position:sticky;top:0;z-index:2;display:grid;grid-template-columns:minmax(420px,1.6fr) 110px minmax(220px,.8fr);gap:10px;padding:10px 12px;background:#f8fafc;border-bottom:1px solid var(--line);font-size:12px;color:#475467;font-weight:700}
.tree-row{display:grid;grid-template-columns:minmax(420px,1.6fr) 110px minmax(220px,.8fr);gap:10px;align-items:center;padding:7px 12px;border-bottom:1px solid #eef2f7;min-height:40px}.tree-row:hover{background:#f8fbff}.tree-row.selected{background:#eff6ff;box-shadow:inset 3px 0 0 var(--blue)}
.name{display:flex;align-items:center;gap:7px;min-width:0}.label{font-weight:700;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.id{color:#94a3b8;font-size:11px}.twisty{width:22px;height:22px;display:inline-flex;align-items:center;justify-content:center;border:1px solid #d9e2ec;border-radius:6px;background:#fff;color:#475467}.twisty.empty{visibility:hidden}.indent{display:inline-block;width:calc(var(--level) * 17px);flex:0 0 calc(var(--level) * 17px);height:20px;border-left:1px solid rgba(148,163,184,.25)}
.kind{font-size:11px;border-radius:999px;padding:3px 7px;background:#eef2f7}.kind.concept{color:#be123c;background:#ffe4e6}.kind.entity{color:#1d4ed8;background:#dbeafe}.kind.group{color:#5b21b6;background:#ede9fe}.category,.relation{font-size:12px;color:#667085;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}.section-title{font-size:13px;margin:0 0 8px;color:#344054}.kv{font-size:13px;line-height:1.65;margin:8px 0}.kv b{color:#344054}.empty{padding:26px;text-align:center;color:var(--muted)}
.rel{border-left:3px solid var(--green);background:#f8fafc;border-radius:6px;padding:7px 9px;margin:6px 0;font-size:12px}.rel b{color:#0f766e}.detail-title{font-size:18px;margin:0 0 4px}.countbar{font-size:12px;color:#667085;margin:0 0 8px}
@media(max-width:1160px){main{grid-template-columns:1fr;height:auto}aside,#details{border:0}.tree-head,.tree-row{grid-template-columns:1fr}.category{display:none}}
</style>
</head>
<body>
<header>
  <div><h1>%%TITLE%%</h1><div class="sub">概念树和实体分面统一展示；点击任意行查看定义、证据、别名与邻近关系。</div></div>
  <div class="toolbar"><input id="search" placeholder="搜索名称 / ID / 类别 / 定义" style="width:300px"/><button id="expandAll">展开全部</button><button id="collapseAll">折叠到一级</button></div>
</header>
<main>
  <aside>
    <div class="card"><h2 class="section-title">总体</h2><div class="stats"><div class="stat"><b id="conceptCount"></b><span class="muted">概念</span></div><div class="stat"><b id="entityCount"></b><span class="muted">实体</span></div><div class="stat"><b id="rootCount"></b><span class="muted">根/分面</span></div><div class="stat"><b id="edgeCount"></b><span class="muted">关系</span></div></div></div>
    <div class="card"><h2 class="section-title">显示</h2><label><input type="checkbox" id="showConcept" checked/> 显示概念</label><br/><label><input type="checkbox" id="showEntity" checked/> 显示实体</label><br/><label><input type="checkbox" id="showGroup" checked/> 显示分组</label><div class="legend"><span><i class="dot" style="background:var(--red)"></i>概念</span><span><i class="dot" style="background:var(--blue)"></i>实体</span><span><i class="dot" style="background:var(--violet)"></i>分组</span></div></div>
    <div class="card"><h2 class="section-title">当前视图</h2><div id="visibleInfo" class="countbar"></div><div class="muted">搜索会保留命中节点及其上级路径。实体统一放在“实体分面”中，避免同一实体在多个疾病概念下重复出现。</div></div>
  </aside>
  <section class="center"><div class="tree-wrap"><div class="tree-head"><span>名称</span><span>类型</span><span>类别/关系</span></div><div id="treeRows"></div></div></section>
  <aside id="details"><div class="card"><h2 class="section-title">节点详情</h2><div id="info" class="muted">点击中间任意一行。</div></div><div class="card"><h2 class="section-title">邻近结构</h2><div id="neighbors" class="muted">选择节点后显示。</div></div></aside>
</main>
<script>
const payload=%%GRAPH%%;
const info=document.getElementById('info'), neighbors=document.getElementById('neighbors'), treeRows=document.getElementById('treeRows'), search=document.getElementById('search');
let selectedRow=''; const rows=payload.rows||[]; let collapsed=new Set(rows.filter(r=>r.depth>=1).map(r=>r.row_id));
const byRow=new Map(rows.map(n=>[n.row_id,n])), children=new Map(), parentOf=new Map(), edgesByNode=new Map();
rows.forEach(r=>{if(r.parent){parentOf.set(r.row_id,r.parent); if(!children.has(r.parent))children.set(r.parent,[]); children.get(r.parent).push(r.row_id)}});
(payload.edges||[]).forEach(e=>{[e.source,e.target].forEach(id=>{if(!edgesByNode.has(id))edgesByNode.set(id,[]); edgesByNode.get(id).push(e)})});
document.getElementById('conceptCount').textContent=payload.concept_count||0; document.getElementById('entityCount').textContent=payload.entity_count||0; document.getElementById('rootCount').textContent=payload.root_count||0; document.getElementById('edgeCount').textContent=payload.edge_count||0;
function allowed(n){if(n.kind==='Concept'&&!document.getElementById('showConcept').checked)return false;if(n.kind==='Entity'&&!document.getElementById('showEntity').checked)return false;if(n.kind==='Group'&&!document.getElementById('showGroup').checked)return false;return true}
function isHiddenByCollapse(n){let p=n.parent;while(p){if(collapsed.has(p))return true;p=parentOf.get(p)}return false}
function matches(n,q){return !q||(`${n.label} ${n.id} ${n.category||''} ${n.definition||''} ${n.relation||''} ${n.aliases||''} ${n.evidence||''}`).toLowerCase().includes(q)}
function ancestorIds(rowId){const out=new Set();let p=parentOf.get(rowId);while(p){out.add(p);p=parentOf.get(p)}return out}
function render(){
  const q=search.value.trim().toLowerCase(); const showConcept=document.getElementById('showConcept').checked; const showEntity=document.getElementById('showEntity').checked; let keep=new Set();
  if(q){rows.forEach(n=>{if(matches(n,q)){keep.add(n.row_id);ancestorIds(n.row_id).forEach(x=>keep.add(x))}})}
  const shown=rows.filter(n=>{if(!allowed(n))return false;if(q&&!keep.has(n.row_id))return false;if(!q&&isHiddenByCollapse(n))return false;return true});
  document.getElementById('visibleInfo').textContent=`显示 ${shown.length} / ${rows.length} 行`;
  treeRows.innerHTML=shown.length?shown.map(rowHtml).join(''):'<div class="empty">没有匹配的节点。</div>';
}
function rowHtml(n){
  const hasKids=(children.get(n.row_id)||[]).length>0; const kind=n.kind==='Concept'?'concept':(n.kind==='Group'?'group':'entity'); const rel=n.relation||n.category||''; const button=hasKids?`<button class="twisty" data-toggle="${escapeAttr(n.row_id)}">${collapsed.has(n.row_id)?'+' : '-'}</button>`:'<span class="twisty empty"></span>';
  const label=n.kind==='Group'?n.label:`${n.label}`;
  return `<div class="tree-row ${selectedRow===n.row_id?'selected':''}" data-row="${escapeAttr(n.row_id)}" style="--level:${Math.min(n.depth,10)}"><div class="name"><span class="indent"></span>${button}<span><span class="label">${escapeHtml(label)}</span><br><span class="id">${escapeHtml(n.id)}</span></span></div><div><span class="kind ${kind}">${n.kind==='Concept'?'概念':(n.kind==='Group'?'分组':'实体')}</span></div><div class="category">${escapeHtml(rel)}</div></div>`;
}
treeRows.addEventListener('click',evt=>{const toggle=evt.target.closest('[data-toggle]');if(toggle){evt.stopPropagation();const id=toggle.dataset.toggle;collapsed.has(id)?collapsed.delete(id):collapsed.add(id);render();return}const row=evt.target.closest('[data-row]');if(row)show(row.dataset.row)});
function show(rowId){
  const n=byRow.get(rowId); if(!n)return; selectedRow=rowId; const lineage=[]; let p=n.parent; while(p){const item=byRow.get(p); if(item)lineage.unshift(item);p=parentOf.get(p)} const kids=(children.get(rowId)||[]).map(cid=>byRow.get(cid)).filter(Boolean);
  const aliases=n.aliases?`<div class="kv"><b>同义词：</b>${escapeHtml(n.aliases)}</div>`:''; const status=n.status?`<div class="kv"><b>状态：</b>${escapeHtml(n.status)}</div>`:''; const docs=n.source_docs?`<div class="kv"><b>来源文档：</b>${escapeHtml(n.source_docs)}</div>`:''; const evidence=n.evidence?`<div class="kv"><b>证据：</b>${escapeHtml(n.evidence)}</div>`:'';
  info.innerHTML=`<h3 class="detail-title">${escapeHtml(n.label)}</h3><div class="muted">${escapeHtml(n.id)}</div><div class="kv"><b>类型：</b>${n.kind==='Concept'?'概念 Concept':(n.kind==='Group'?'导航分组':'实体 Entity')}</div><div class="kv"><b>类别/关系：</b>${escapeHtml(n.relation||n.category||'')}</div><div class="kv"><b>定义：</b>${escapeHtml(n.definition||'')}</div>${aliases}${status}${docs}${evidence}<div class="kv"><b>路径：</b><br>${lineage.map(x=>`<span class="pill">${escapeHtml(x.label)}</span>`).join(' ')||'<span class="muted">根节点</span>'}</div>`;
  const related=(edgesByNode.get(n.id)||[]).slice(0,80);
  neighbors.innerHTML=`<div class="kv"><b>下级节点：</b><br>${kids.slice(0,80).map(x=>`<span class="pill">${escapeHtml(x.label)}</span>`).join(' ')||'<span class="muted">无</span>'}</div><div class="kv"><b>相关关系：</b>${related.length?related.map(edgeHtml).join(''):'<div class="muted">无直接关系</div>'}</div>`; render();
}
function edgeHtml(e){return `<div class="rel">${escapeHtml(e.source_label)} <b>→ ${escapeHtml(e.relation)} →</b> ${escapeHtml(e.target_label)}${e.evidence?`<br><span class="muted">${escapeHtml(e.evidence)}</span>`:''}</div>`}
document.getElementById('expandAll').onclick=()=>{collapsed.clear();render()}; document.getElementById('collapseAll').onclick=()=>{collapsed=new Set(rows.filter(n=>n.depth>=1&&(children.get(n.row_id)||[]).length).map(n=>n.row_id));render()}; document.getElementById('showConcept').onchange=render; document.getElementById('showEntity').onchange=render; document.getElementById('showGroup').onchange=render; search.oninput=render;
function escapeHtml(s){return String(s??'').replace(/[&<>"']/g,m=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[m]))}
function escapeAttr(s){return escapeHtml(s)}
render();
</script>
</body>
</html>"""
    template = template.replace("%%TITLE%%", html.escape(title)).replace("%%GRAPH%%", json.dumps(payload, ensure_ascii=False))
    Path(path).write_text(template, encoding="utf-8")


def write_main_style_html(nodes, edges, path):
    """Render merged KG with the same friendly browser used by the main KG."""
    try:
        root_dir = Path(__file__).resolve().parents[1]
        if str(root_dir) not in sys.path:
            sys.path.insert(0, str(root_dir))
        from build_colorectal_kg import write_friendly_html

        normalized_nodes = {}
        for node_id, node in nodes.items():
            category = node.get("category") or "Other Entity"
            normalized_nodes[node_id] = {
                "id": node.get("id") or node_id,
                "label": node.get("label") or node_id,
                "node_kind": node.get("node_kind") or "Entity",
                "in_colorectal_tree": node.get("in_colorectal_tree", "N"),
                "category": category,
                "research_axis": node.get("research_axis") or CATEGORY_TO_AXIS.get(category, "Other"),
                "semantic_types": node.get("semantic_types", ""),
                "depth_from_root": node.get("depth_from_root", ""),
                "active": node.get("active", ""),
                "concept_status": node.get("concept_status", ""),
                "umls_cui": node.get("umls_cui", ""),
                "icd_o_3_code": node.get("icd_o_3_code", ""),
                "definition": node.get("definition") or node.get("evidence", ""),
                "aliases": node.get("aliases") or node.get("label") or node_id,
                "alias_count": node.get("alias_count", ""),
                "nci_url": node.get("nci_url", ""),
                "kg_status": node.get("kg_status", ""),
                "source_docs": node.get("source_docs", ""),
                "evidence": node.get("evidence", ""),
            }

        normalized_edges = []
        for edge in edges:
            if edge.get("source") not in normalized_nodes or edge.get("target") not in normalized_nodes:
                continue
            normalized_edges.append(
                {
                    "source": edge.get("source", ""),
                    "source_label": edge.get("source_label", normalized_nodes[edge.get("source", "")]["label"]),
                    "relation_code": edge.get("relation_code", ""),
                    "relation_type": edge.get("relation_type", "RELATED_TO"),
                    "relation_label_cn": edge.get("relation_label_cn") or edge.get("relation_type", "相关"),
                    "target": edge.get("target", ""),
                    "target_label": edge.get("target_label", normalized_nodes[edge.get("target", "")]["label"]),
                    "relation_group": edge.get("relation_group", "incremental"),
                    "direction": edge.get("direction", "out"),
                    "research_tier": edge.get("research_tier", "supporting"),
                    "source_docs": edge.get("source_docs", ""),
                    "evidence": edge.get("evidence", ""),
                }
            )
        write_friendly_html(normalized_nodes, normalized_edges, path)
    except Exception as exc:
        fallback_title = f"更新后的结直肠癌知识图谱（主样式渲染失败，已回退：{exc}）"
        write_mini_html(nodes, edges[:12000], path, fallback_title)


def write_mini_html(nodes, edges, path, title):
    node_list = list(nodes.values())
    edge_list = list(edges)
    degree = Counter()
    for edge in edge_list:
        degree[edge["source"]] += 1
        degree[edge["target"]] += 1
    graph = {
        "nodes": [
            {
                "id": node["id"],
                "label": node["label"],
                "category": node.get("category", ""),
                "kind": node.get("node_kind", ""),
                "status": node.get("kg_status", ""),
                "definition": node.get("definition", ""),
                "evidence": node.get("evidence", ""),
                "degree": degree[node["id"]],
            }
            for node in node_list
        ],
        "edges": [
            {
                "source": edge["source"],
                "target": edge["target"],
                "label": edge.get("relation_label_cn") or edge.get("relation_type"),
                "type": edge.get("relation_type", ""),
            }
            for edge in edge_list
        ],
    }
    template = """<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"/><meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>%%TITLE%%</title>
<style>
body{margin:0;font-family:Arial,"Microsoft YaHei",sans-serif;background:#f6f8fb;color:#172033}
header{padding:16px 20px;background:white;border-bottom:1px solid #d9e2ec;display:flex;gap:16px;align-items:center;justify-content:space-between}
h1{font-size:20px;margin:0}.muted{color:#667085;font-size:13px}
main{display:grid;grid-template-columns:1fr 340px;height:calc(100vh - 68px)}
canvas{width:100%;height:100%;display:block}aside{background:white;border-left:1px solid #d9e2ec;padding:16px;overflow:auto}
input{width:260px;border:1px solid #cbd5e1;border-radius:8px;padding:9px}
.pill{display:inline-block;background:#eef2f7;border-radius:999px;padding:3px 7px;margin:3px;font-size:11px}
.card{border:1px solid #d9e2ec;border-radius:8px;padding:12px;background:#fff}
</style></head><body>
<header><div><h1>%%TITLE%%</h1><div class="muted">节点是概念/实体；边表示资料或图谱中的关系。</div></div><input id="search" placeholder="搜索节点"/></header>
<main><canvas id="c"></canvas><aside><h3>节点详情</h3><div id="info" class="card muted">点击节点查看。</div></aside></main>
<script>
const graph=%%GRAPH%%; const canvas=document.getElementById('c'); const ctx=canvas.getContext('2d'); const info=document.getElementById('info');
let w=0,h=0,scale=1,ox=0,oy=0,selected=null,drag=null,pan=false,last=null;
const byId=new Map(graph.nodes.map(n=>[n.id,n]));
function resize(){const r=canvas.getBoundingClientRect(),d=devicePixelRatio||1;canvas.width=r.width*d;canvas.height=r.height*d;ctx.setTransform(d,0,0,d,0,0);w=r.width;h=r.height;layout();draw();}
function layout(){const cx=w/2,cy=h/2;graph.nodes.sort((a,b)=>b.degree-a.degree);graph.nodes.forEach((n,i)=>{if(n._moved)return;const a=i*2.399963;const r=40+Math.sqrt(i)*25;n.x=cx+Math.cos(a)*r;n.y=cy+Math.sin(a)*r;});}
function color(n){if(n.status&&n.status.includes('new'))return '#f97316';if(n.kind==='Concept')return '#e11d48';if(n.category&&n.category.includes('Gene'))return '#0f766e';if(n.category&&n.category.includes('Treatment'))return '#be123c';return '#2563eb';}
function pass(n){const q=document.getElementById('search').value.trim().toLowerCase();return !q||(n.id+' '+n.label+' '+n.category).toLowerCase().includes(q);}
function visible(){const ids=new Set(graph.nodes.filter(pass).map(n=>n.id));return {nodes:graph.nodes.filter(n=>ids.has(n.id)),edges:graph.edges.filter(e=>ids.has(e.source)&&ids.has(e.target))};}
function draw(){const v=visible();ctx.clearRect(0,0,w,h);ctx.save();ctx.translate(ox,oy);ctx.scale(scale,scale);v.edges.forEach(e=>{const s=byId.get(e.source),t=byId.get(e.target);ctx.strokeStyle='rgba(71,85,105,.22)';ctx.beginPath();ctx.moveTo(s.x,s.y);ctx.lineTo(t.x,t.y);ctx.stroke();});v.nodes.forEach(n=>{const r=n.id==='C2955'?15:Math.max(6,Math.min(13,5+Math.sqrt(n.degree||1)));ctx.beginPath();ctx.fillStyle=color(n);ctx.strokeStyle=selected&&selected.id===n.id?'#111827':'#fff';ctx.lineWidth=selected&&selected.id===n.id?3/scale:1.5/scale;ctx.arc(n.x,n.y,r,0,Math.PI*2);ctx.fill();ctx.stroke();if(scale>.55||selected&&selected.id===n.id){ctx.fillStyle='#172033';ctx.font=`${12/scale}px Arial`;ctx.fillText(n.label.slice(0,42),n.x+r+4,n.y+4);}});ctx.restore();}
function world(e){const r=canvas.getBoundingClientRect();return{x:(e.clientX-r.left-ox)/scale,y:(e.clientY-r.top-oy)/scale};}
function hit(p){let b=null,bd=18/scale;visible().nodes.forEach(n=>{const d=Math.hypot(n.x-p.x,n.y-p.y);if(d<bd){b=n;bd=d;}});return b;}
canvas.onmousedown=e=>{const p=world(e),n=hit(p);last={x:e.clientX,y:e.clientY};if(n){selected=n;drag=n;show(n);}else pan=true;draw();}
canvas.onmousemove=e=>{if(drag){const p=world(e);drag.x=p.x;drag.y=p.y;drag._moved=true;draw();}else if(pan){ox+=e.clientX-last.x;oy+=e.clientY-last.y;last={x:e.clientX,y:e.clientY};draw();}}
onmouseup=()=>{drag=null;pan=false};canvas.onwheel=e=>{e.preventDefault();const old=scale;scale*=e.deltaY<0?1.12:.89;scale=Math.max(.25,Math.min(4,scale));draw();}
function esc(s){return String(s||'').replace(/[&<>"']/g,m=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[m]));}
function show(n){info.innerHTML=`<b>${esc(n.label)}</b><br><span class="muted">${esc(n.id)} · ${esc(n.kind)} · ${esc(n.category)}</span><p>${esc(n.definition||'')}</p><p><b>状态</b>：${esc(n.status||'existing')}</p><p><b>证据</b>：${esc(n.evidence||'')}</p>`}
document.getElementById('search').oninput=draw;addEventListener('resize',resize);resize();ox=0;oy=0;draw();
</script></body></html>"""
    template = template.replace("%%TITLE%%", html.escape(title)).replace("%%GRAPH%%", json.dumps(graph, ensure_ascii=False))
    Path(path).write_text(template, encoding="utf-8")


def cmd_extract(args):
    config = load_config(args.config)
    output = Path(args.output) if args.output else DEFAULT_OUT_DIR / time.strftime("mini_%Y%m%d_%H%M%S")
    report = build_mini_kg(args.input, args.kg_dir, output, config=config, use_llm=args.use_llm)
    print(json.dumps({"output_dir": str(output.resolve()), **report}, ensure_ascii=False, indent=2))


def cmd_merge(args):
    report = merge_kg(args.base_kg, args.mini, args.output, accept_candidates=args.accept_candidates)
    print(json.dumps(report, ensure_ascii=False, indent=2))


def cmd_ask(args):
    config = load_config(args.config)
    result = answer_question(
        args.question,
        args.kg_dir,
        config=config,
        top_k=args.top_k,
        use_external=args.external,
        external_k=args.external_k,
        answer_mode=args.answer_mode,
    )
    if args.output:
        Path(args.output).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


def build_parser():
    parser = argparse.ArgumentParser(description="结直肠癌知识图谱增量抽取、合并与问答")
    sub = parser.add_subparsers(required=True)

    p = sub.add_parser("extract", help="从新资料生成小知识图谱")
    p.add_argument("--input", required=True, help="新资料文件或目录")
    p.add_argument("--kg-dir", default=str(BASE_KG_DIR), help="基础知识图谱目录")
    p.add_argument("--output", help="小图谱输出目录")
    p.add_argument("--config", help="API 配置 JSON")
    p.add_argument("--use-llm", action="store_true", help="启用配置中的大模型抽取")
    p.set_defaults(func=cmd_extract)

    p = sub.add_parser("merge", help="把小图谱合并进主图谱")
    p.add_argument("--base-kg", default=str(BASE_KG_DIR), help="基础知识图谱目录")
    p.add_argument("--mini", required=True, help="mini 图谱目录")
    p.add_argument("--output", required=True, help="更新后图谱输出目录")
    p.add_argument("--accept-candidates", action="store_true", help="接受新增候选实体并写入更新后图谱")
    p.set_defaults(func=cmd_merge)

    p = sub.add_parser("ask", help="基于图谱问答")
    p.add_argument("--question", required=True, help="问题")
    p.add_argument("--kg-dir", default=str(BASE_KG_DIR), help="知识图谱目录")
    p.add_argument("--config", help="API 配置 JSON；不提供则只返回检索上下文")
    p.add_argument("--top-k", type=int, default=20)
    p.add_argument("--external", action="store_true", help="同时检索 PubMed 第三方资料")
    p.add_argument("--external-k", type=int, default=5, help="PubMed 返回数量")
    p.add_argument(
        "--answer-mode",
        choices=["graph", "literature", "comprehensive"],
        default="graph",
        help="graph=图谱优先；literature=PubMed文献综述；comprehensive=图谱+PubMed+AI综合知识",
    )
    p.add_argument("--output", help="答案 JSON 输出路径")
    p.set_defaults(func=cmd_ask)
    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
